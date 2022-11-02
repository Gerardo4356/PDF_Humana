from copy import copy
from importlib.metadata import FileHash
import math
from tkinter.tix import InputOnly
from unittest import skip
from numpy import append
from pdfminer.high_level import extract_text #pip install pdfminer.six
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls  # para color de tabla
from docx.oxml import parse_xml  # para color de tabla
# Agregar referencia de tabla
# https://stackoverflow.com/questions/54514666/cross-referenceable-figure-numbers-by-section-with-python-docx/54534731#54534731
from docx.oxml import OxmlElement  # Para caption (referencias de tabla)
from docx.oxml.ns import qn  # Para caption (referencias de tabla)
from docx.enum.text import WD_COLOR_INDEX  # Subrayar
from docx.shared import RGBColor  # Color letra
from docx.shared import Pt #Tamaño de letra
from docx.shared import Cm, Inches # Ancho de columnas de tabla
from docx.enum.text import WD_COLOR_INDEX #Subrayar

from datetime import date
from datetime import datetime
from datetime import timedelta
from dateutil.relativedelta import relativedelta
import calendar

import os
import pyperclip as pc #para copiar al clipboard
os.system("cls")


def edad(cumple):
    today = date.today()
    edad = today.year - cumple.year - \
        ((today.month, today.day) < (cumple.month, cumple.day))
    return edad

def diferencia_fechas(strfecha, strfecha2):
    fecha1 = datetime.strptime(strfecha, '%d/%m/%Y')
    fecha2 = datetime.strptime(strfecha2, '%d/%m/%Y')
    diferencia = fecha2 - fecha1
    diferencia = diferencia.days
    return diferencia

def restar_dias(fecha, dias):
    resta = datetime.strptime(fecha, '%d/%m/%Y') + timedelta(days=dias)
    resta = resta.strftime("%d/%m/%Y") # Para que no cambie el formato de la fecha
    return str(resta)

def copiar(array):
    x = str(array).replace("[","")
    x = x.replace("]","")
    x = x.replace("'","")
    x = x.replace(", ","\n")
    pc.copy(x)
    
def diagnostico(path="PDF/test0.pdf"):
    from datetime import date
    text = extract_text(path).split("\n")


    #region Extrayendo texto del pdf 
    # Obtener datos del encabezado
    pension_minima = "5,836"  # 2022 lo da el gobierno cada año
    semanas_cotizadas = text[23]
    semanas_reconocidas = text[34]
    semanas_desconocidas = text[36]
    semanas_reintegradas = text[38]
    
    # Obteniendo datos
    NOMBRE = text[6]
    NSS = text[10]
    CURP = text[14]
    cumple_anio = "19"+CURP[4:6]
    cumple_mes = CURP[6:8].lstrip("0")
    cumple_dia = CURP[8:10].lstrip("0")
    EDAD = edad(date(int(cumple_anio), int(cumple_mes), int(cumple_dia)))


    #Obteniendo datos linea por línea

        # For para concontrar la columna de fechas de movimiento
        # Patrones es la cantidad de patrones que ha tenido
        # Fechas es una variable candado para saltarnos la fecha de alta y de baja (que no es la de fecha de movimiento)
        # En fechas_de_movimiento se guardan todas las fechas de la columna fecha de movimiento
        # salarios es variable candado para saltarnos el del encabezado
    
    patrones = 0
    lastpatrones = 0
    fechas = 1
    salarios = 0
    fechas_movimiento = []
    salarios_base = []
    vigente = False
    for i, line in enumerate(text):
        if line == "":
            continue

        if "Salario Base de Cotización */" == line:
            patrones += 1
            fechas = 1
            salarios = 0
        if "Salario Base" == line:
            salarios = 1

        if "Fecha de alta" == line:
            fechas = 0

        if "Vigente" == line:
            # Detecto un vigente, y guardo una varible como true, para saltarme el cambio de patrón
            if date.today().day >= 30:
                ultimo_dia_mes = str(calendar.monthrange(date.today().year,date.today().month+1)[1]) + "/" + str(date.today().month+1) + "/" + str(date.today().year)
            else:
                ultimo_dia_mes = str(calendar.monthrange(date.today().year,date.today().month)[1]) + "/" + str(date.today().month) + "/" + str(date.today().year)
 
            
            fechas_movimiento.append("")
            salarios_base.append("")
            fechas_movimiento.append(ultimo_dia_mes)
            salarios_base.append("0")
            vigente = True

        if patrones > 0 and fechas == 1:
            try:
                date = datetime.strptime(line, '%d/%m/%Y')
                fechas_movimiento.append(line)
                vigente = False
            except:
                pass
        if patrones > 0 and salarios == 1:
            if '$' in line:
                salarios_base.append(line)
                vigente = False

        if lastpatrones != patrones:
            # Hubo cambio de patron
            if not vigente:
                fechas_movimiento.append("")
                salarios_base.append("")
        lastpatrones = patrones
    # Quitar signo de pesos en salarios_base
    for i,salarios in enumerate(salarios_base): salarios_base[i] = salarios_base[i].replace("$","")

    #endregion extrayendo texto del pdf

    """ Cambio con relación a la versión anterior: Antes se cortaba desde el inicio, hasta donde llegaran 250 semanas cotizadas. Ahora se elimina esta limitante, y se itera hasta que termina todas las fechas. """


    #region LASTI Puede dar un error en el futuro con el doble patrón.
    dias = []
    dias_laborados = 0
    lasti = 0
    for i,fecha in enumerate(fechas_movimiento):
        if dias_laborados <= 1750:
            if fecha != "" and fechas_movimiento[i-1] != "":
                diff = diferencia_fechas(fecha,fechas_movimiento[i-1])
                if diff < 1: 
                    dias.append("") 
                    # input("Patrón sin valor")
                    print("Patrón sin valor")
                else: #Si hay diferencia positiva de días, se agregan a los arrays
                    # print(fechas_movimiento[i-1])
                    # print(fecha)
                    dias.append(diff)
                    dias_laborados += diff
                    lasti = i
            else:
                dias.append("")
        else: continue
    #endregion LASTI

    #region Doble Patron
    #Checar si existe doble patrón
    doble_patron = False
    for i,fecha in enumerate(fechas_movimiento):
        if i > 1 and i < lasti+1:  #A partir de esta fila (2) es cuando se puede detectar el doble patrón. 
            if fechas_movimiento[i-1] == "": #Si hubo un cambio de patrón, se debe verificar si hubo doble patrón
                if fecha != "" and fechas_movimiento[i-2] != "":    # Cuando haya patrones sin valor, se skipea
                    if diferencia_fechas(fecha,fechas_movimiento[i-2]) < 0:#Se resta la fecha de baja actual, contra la fecha de alta del patrón anterior, si da número negativo, quiere decir que interfiere en las fechas, por lo tanto, trabajo con dos patrones en la misma fecha
                        doble_patron = True




    # Corregir doble patrón
    ##################
    # Nota: Reciclar el proceso anterior, pero eliminando el corte de 250 semanas
    ##################
    while doble_patron:
        #region Proceso doble patron
        print("!! Doble patrón")
        id_patrones = []
        tipo = []       
        contador = 0    
        #Aqui se escribe dos arrays, uno para id de patrones, y otro para starus
        for i,fecha in enumerate(fechas_movimiento):
            if fecha != "":
                id_patrones.append(contador)
            else:
                tipo.append("")
                id_patrones.append("")
                contador = contador + 1
            #Para escribir los tipos
            if i<lasti:
                if i>0 and fecha != "" and fechas_movimiento[i-1] == "": tipo.append("BAJA")
                if i>0 and fecha != ""                    and fechas_movimiento[i-1] != "" and fechas_movimiento[i+1] != "": tipo.append("CAMBIO")
                if i>0 and fecha != "" and fechas_movimiento[i-1] != "" and fechas_movimiento[i+1] == "": tipo.append("ALTA")
       
        tipo[lasti] = "ALTA" #Puede que la última fecha sea de un cambio, pero para el corte, lo convertimos en alta.

        rango_inicial = 0
        rango_final = 0
        #Sacar rango inicial y final para saber donde vamos a cambiar el array original
        for i,fecha in enumerate(fechas_movimiento):
            # if i > 1 and i< lasti+1:  #A partir de esta fila (2) es cuando se puede detectar el doble patrón. También si ya cruzamos el último antes de las 250 semanas, se cancela el chequeo
            if i > 1 and i< len(fechas_movimiento):  #A partir de esta fila (2) es cuando se puede detectar el doble patrón. 
                if fechas_movimiento[i-1] == "": #Si hubo un cambio de patrón, se debe verificar si hubo doble patrón
                    if fecha != "" and fechas_movimiento[i-2] != "":    # Cuando haya patrones sin valor, se skipea
                        if diferencia_fechas(fecha,fechas_movimiento[i-2]) < 0:#Se resta la fecha de baja actual, contra la fecha de alta del patrón anterior, si da número negativo, quiere decir que interfiere en las fechas, por lo tanto, trabajo con dos patrones en la misma fecha
                            if rango_inicial == 0:  #Si no existe un rango inicial, anotarlo
                                rango_inicial = id_patrones[i]-1
                                # print(diferencia_fechas(fecha,fechas_movimiento[i-2]))
                                print("Rango inicial:\t"+str(rango_inicial)+"\t"+str(diferencia_fechas(fecha,fechas_movimiento[i-2])))
                        else:
                            if rango_inicial != 0 and rango_final == 0: # Si ya se escribio rango inicial, y rango final no existe, escribirlo.
                                rango_final = id_patrones[i]-1 
                                print("Rango final:\t"+str(rango_final)+"\t"+str(diferencia_fechas(fecha,fechas_movimiento[i-2])))

        #Obtener el indice del primer patron duplicado
        indice_primer_doble_patron = 0
        indice_ultimo_doble_patron = 0
        for i,id in enumerate(id_patrones):
            if id == rango_inicial and indice_primer_doble_patron== 0: # Si encontramos el primer id, y no hay primer indice, anotarlo.
                indice_primer_doble_patron = i
            if id == rango_final: #Si encontramos el último id, anotarlo, y se sobreescribirá hasta quedar el último
                indice_ultimo_doble_patron = i
    
        #Crear el contenido cambiado
            #Ocupo comparar la fecha con una muy lejana, para determinar que fecha es más reciente. No puedo poner la fecha de hoy, por que si es vigente, va a entrar en conflicto.
        
        #Extraigo la informacion que requiero, de la parte donde se intersecta el doble patrón
        indice_filtro = []
        for i in range(indice_primer_doble_patron, indice_ultimo_doble_patron + 1): indice_filtro.append(i)
        fechas_filtradas = fechas_movimiento[indice_primer_doble_patron: indice_ultimo_doble_patron+1]
        sueldos_filtrados = salarios_base[indice_primer_doble_patron: indice_ultimo_doble_patron+1]
        id_filtrados = id_patrones[indice_primer_doble_patron: indice_ultimo_doble_patron+1]
        tipos_filtrados = tipo[indice_primer_doble_patron: indice_ultimo_doble_patron+1]




        fechas_vs_pivote = []
        for i in fechas_filtradas:
            if i != "" : 
                fechas_vs_pivote.append(diferencia_fechas(i,"31/12/2099"))
            else: fechas_vs_pivote.append("")

        # Datos filtrados
        n_indice_filtro = []
        n_fechas_vs_pivote = []
        n_fechas_filtradas = []
        n_sueldos_filtrados = []
        n_id_filtrados = []
        n_tipos_filtrados = []

        # Elimino los espacios donde hubo separacion entre patrones
        for i,j in enumerate(fechas_filtradas):
            if j != "":
                n_indice_filtro.append(indice_filtro[i])
                n_fechas_vs_pivote.append(fechas_vs_pivote[i])
                n_fechas_filtradas.append(fechas_filtradas[i])
                n_sueldos_filtrados.append(sueldos_filtrados[i])
                n_id_filtrados.append(id_filtrados[i])
                n_tipos_filtrados.append(tipos_filtrados[i])



        print("------")
        
        #Creo un combinado (ya ordenado de )        
        filtro = sorted(zip(n_fechas_vs_pivote, n_indice_filtro, n_id_filtrados, n_tipos_filtrados, n_fechas_filtradas, n_sueldos_filtrados))

        saldo_acumulado = []
        for i in range(len(filtro)-1,-1,-1): #For en reversa, para analizar según el dado de alta
            if filtro[i][3] == 'ALTA':
                saldo_acumulado.append((filtro[i][2],filtro[i][5]))
            if filtro[i][3] == 'BAJA':
                #Para eliminar el patron que ya se dio de baja. Si el númeroo de patrón coincide con columna 0 de saldo acumulado, anotar indice, y eliminarlo. (El indice se anoto en j)
                for j in range(len(saldo_acumulado)):
                    if saldo_acumulado[j][0] == filtro[i][2]:
                        index_cambio = j
                saldo_acumulado.remove(saldo_acumulado[index_cambio])

            if filtro[i][3] == 'CAMBIO':
                
                for j in range(len(saldo_acumulado)):
                    if saldo_acumulado[j][0] == filtro[i][2]:
                        index_cambio = j
                saldo_acumulado[index_cambio] = (filtro[i][2],filtro[i][5])

            saldo_nuevo = 0
            for _,saldo in saldo_acumulado:
                saldo_nuevo = saldo_nuevo+float(saldo)
            filtro[i] = (filtro[i][0], filtro[i][1], filtro[i][2], filtro[i][3], filtro[i][4], saldo_nuevo)
        
        #Ordenar el documento
        fechas_filtradas_ordenadas = []
        for i in filtro: fechas_filtradas_ordenadas.append(i[4])
        
        saldos_filtradas_ordenadas = []
        for i in filtro: saldos_filtradas_ordenadas.append(i[5])
        
        fechas_movimiento_nuevo = fechas_movimiento[:indice_primer_doble_patron] + fechas_filtradas_ordenadas + fechas_movimiento[indice_ultimo_doble_patron:]
        salarios_base_nuevo = salarios_base[:indice_primer_doble_patron] + saldos_filtradas_ordenadas +  salarios_base[indice_ultimo_doble_patron:]
        #region Clipboard temporal
        # # Copiar fechas y salarios solo hasta donde fueron 1750 días
        # x = str(fechas_movimiento[:lasti+1]).replace("[","")
        # x = x.replace("]","")
        # x = x.replace("'","")
        # x = x.replace(", ","\n")
        # pc.copy(x)
        # input("--Fechas de doble patrón en clipboard")
        # x = str(salarios_base[:lasti+1]).replace("[","")
        # x = x.replace("]","")
        # x = x.replace("'","")
        # x = x.replace(", ","\n")
        # pc.copy(x)
        # input("--Salarios de doble patrón en  clipboard")
        # id_patrones = str(id_patrones[:lasti+1]).replace("[","")
        # id_patrones = id_patrones.replace("]","")
        # id_patrones = id_patrones.replace("'","")
        # id_patrones = id_patrones.replace(", ","\n")
        # pc.copy(id_patrones)
        # input("--ID en clipboard")
        # tipo = str(tipo[:lasti+1]).replace("[","")
        # tipo = tipo.replace("]","")
        # tipo = tipo.replace("'","")
        # tipo = tipo.replace(", ","\n")
        # pc.copy(tipo)
        # input("--TIPO en clipboard")
        #endregion Temporal

        #endregion
        
        fechas_movimiento = fechas_movimiento_nuevo
        salarios_base = salarios_base_nuevo





        #region LASTI Puede dar un error en el futuro con el doble patrón.
        dias = []
        dias_laborados = 0
        lasti = 0
        for i,fecha in enumerate(fechas_movimiento):
            if dias_laborados <= 1750:
                if fecha != "" and fechas_movimiento[i-1] != "":
                    diff = diferencia_fechas(fecha,fechas_movimiento[i-1])
                    if diff < 1: 
                        dias.append("") 
                        # input("Patrón sin valor")
                        print("Patrón sin valor")
                    else: #Si hay diferencia positiva de días, se agregan a los arrays
                        # print(fechas_movimiento[i-1])
                        # print(fecha)
                        dias.append(diff)
                        dias_laborados += diff
                        lasti = i
                else:
                    dias.append("")
            else: continue
        #endregion LASTI




        #Verificar que no haya doble doble patrón
        doble_patron = False
        for i,fecha in enumerate(fechas_movimiento):
            if i > 1 and i < lasti+1:  #A partir de esta fila (2) es cuando se puede detectar el doble patrón. 
                if fechas_movimiento[i-1] == "": #Si hubo un cambio de patrón, se debe verificar si hubo doble patrón
                    if fecha != "" and fechas_movimiento[i-2] != "":    # Cuando haya patrones sin valor, se skipea
                        if diferencia_fechas(fecha,fechas_movimiento[i-2]) < 0:#Se resta la fecha de baja actual, contra la fecha de alta del patrón anterior, si da número negativo, quiere decir que interfiere en las fechas, por lo tanto, trabajo con dos patrones en la misma fecha
                            doble_patron = True

        
    #endregion Doble Patron

    #region Cortar solo a cuando hemos superado los 1750 días
    dias = []
    dias_laborados = 0
    lasti = 0
    for i,fecha in enumerate(fechas_movimiento):
        if dias_laborados <= 1750:
            if fecha != "" and fechas_movimiento[i-1] != "":
                diff = diferencia_fechas(fecha,fechas_movimiento[i-1])
                if diff < 1: 
                    dias.append("") 
                    # input("Patrón sin valor")
                    print("Patrón sin valor")
                else: #Si hay diferencia positiva de días, se agregan a los arrays
                    # print(fechas_movimiento[i-1])
                    # print(fecha)
                    dias.append(diff)
                    dias_laborados += diff
                    lasti = i
            else:
                dias.append("")
        else: continue




    #Se restan días para que sumen 1750
    dias_sobrantes = dias_laborados - 1750
    if dias_sobrantes > 0:
        # Cambiar la última fecha para que sume 1750
        fechas_movimiento[lasti] = restar_dias(fechas_movimiento[lasti],dias_sobrantes)
        dias[-1] = dias[-1]-dias_sobrantes
    

    #endregion Cortar solo a cuando hemos superado los 1750 días

    #region Imprimir anormalidades
    #Imprimir si vigente
    if "Vigente" in text: 
        print("-Vigente")
        vigente = True
    #Imprimir si doble patrón
    if doble_patron: print("-Doble patrón")

    #endregion Imprimir anormalidades

    #region Calculos

    #Calculo de semanas
    semanas = []
    for i, num_dias in enumerate(dias): 
        if dias[i] != "":semanas.append(dias[i]/7)
        else: semanas.append("")
        
    
    #Calculo de ponderacion (semanas*salario)
    ponderacion = []
    for i, num_semanas in enumerate(semanas):
        if semanas[i] != "" and salarios_base[i] != "":
            ponderacion.append(float(semanas[i]) * float(salarios_base[i])) 
        else: ponderacion.append("")
        
    sum_ponderacion = 0
    sum_semanas = 0
    for i in semanas:
        if i != "":
            sum_semanas += i
    
    for i in ponderacion:
        if i in ponderacion:
            if i != "":
                sum_ponderacion += i
                
    #endregion Calculos

    #region Clipboard
    temp_excel = []  
    fechas_movimiento = fechas_movimiento[:lasti+1]
    for i in range(len(fechas_movimiento)):
        temp_excel.append(str(fechas_movimiento[i])+ "	" + str(salarios_base[i]))
    excel = str(temp_excel).replace("[","")
    excel = excel.replace("]","")
    excel = excel.replace("'","")
    excel = excel.replace(", ","\n")
    excel = excel.replace(r"\t ","	")
    excel = excel.replace(r"\t","	")
    # input(excel)
    pc.copy(excel)


    #Copiar fechas y salarios solo hasta donde fueron 1750 días
    # fechas_movimiento = str(fechas_movimiento[:lasti+1]).replace("[","")
    # fechas_movimiento = fechas_movimiento.replace("]","")
    # fechas_movimiento = fechas_movimiento.replace("'","")
    # fechas_movimiento = fechas_movimiento.replace(", ","\n")
    # pc.copy(fechas_movimiento)
    # input("FECHAS EN EL CLIPBOARD. (Presione enter para copiar salarios.)")
    # salarios_base = str(salarios_base[:lasti+1]).replace("[","")
    # salarios_base = salarios_base.replace("]","")
    # salarios_base = salarios_base.replace("'","")
    # salarios_base = salarios_base.replace(", ","\n")
    # pc.copy(salarios_base)
    # input("SALARIOS EN EL CLIPBOARD. (Presione enter para copiar diferenciadias.)")
    # dias = str(dias).replace("[","")
    # dias = dias.replace("]","")
    # dias = dias.replace("'","")
    # dias = dias.replace(", ","\n")
    # pc.copy(dias)
    # input("DIFERENCIA DE DIAS EN CLIPBOARD. (Presione enter para copiar semanas.)")
    # semanas = str(semanas).replace("[","")
    # semanas = semanas.replace("]","")
    # semanas = semanas.replace("'","")
    # semanas = semanas.replace(", ","\n")
    # pc.copy(semanas)
    # input("CANTIDAD DE SEMANAS EN EL CLIPBOARD")
    # ponderacion = str(ponderacion).replace("[","")
    # ponderacion = ponderacion.replace("]","")
    # ponderacion = ponderacion.replace("'","")
    # ponderacion = ponderacion.replace(", ","\n")
    # pc.copy(ponderacion)
    # input("PONDERACION EN EL CLIPBOARD")


    #endregion Clipboard

    #region Tabla
    salario_prom = round(sum_ponderacion/sum_semanas,2)

    tablaVSMi = [0, 1.01, 1.26, 1.51, 1.76, 2.01, 2.26, 2.51, 2.76, 3.01, 3.26, 3.51, 3.76, 4.01, 4.26, 4.51, 4.76, 5.01, 5.26, 5.51, 5.76, 6.01]
    tablaVSMx = [1, 1.25, 1.50, 1.75, 2.00, 2.25, 2.50, 2.75, 3.00, 3.25, 3.50, 3.75, 4.00, 4.25, 4.50, 4.75, 5.00, 5.25, 5.50, 5.75, 6.00, 25.00]
    tablaCB  = [.8000, .7711, .5818, .4923, .4267, .3765, .3368, .3048, .2783, .2560, .2370, .2207, .2065, .1939, .1829, .1730, .1641, .1561, .1488, .1422, .1362, .1300]
    tablaIncr = [.005630, .008140, .01178, .01430, .01615, .01756, .01868, .01958, .02033, .02096, .02149, .02195, .02235, .02271, .02302, .02330, .02355, .02377, .02398, .02416, .02433, .02450]

    tablaEdades = [55, 56, 57, 58, 59, 60, 61, 62, 63, 64, 65]
    tablaPorc = [.75, .75, .75, .75, .75, .75, .80, .85, .90, .95, 1]



    
    
    tabla = [0,1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17]
    tabla[0] = semanas_cotizadas
    tabla[1] = salario_prom
    tabla[2] = ""
    tabla[3] = 1#input("CÓNYUGE (1 o 0): ")
    tabla[4] = 0#input("HIJOS (MENORES O ESTUDIANDO): ")
    tabla[5] = 0#input("PADRES (SOLO FALTA DE ESPOSA E HIJOS): ")
    # input("EL SALARIO MINIMO DF ESTÁ ESTABLECIDO EN 96.22")
    tabla[6] = 96.22
    tabla[7] = float(tabla[1])/float(tabla[6])
    tabla[7] = round(tabla[7],2)
    
    for i, VSMi in enumerate(tablaVSMi):
        if tablaVSMi[i] <= tabla[7] and tablaVSMx[i] >= tabla[7]:
            tabla[8] = tablaCB[i] #Cuantía básica
            tabla[9] = tablaIncr[i] #Incremento anual     
    tabla[10] = tabla[8] * tabla[1] * 365 #Cuantia básica anual en pesos
    
    semanas_excedentes = float(tabla[0]) - 500
    anios_excedentes = math.trunc(semanas_excedentes/52)
    semanas_enteras_excedentes = anios_excedentes * 52
    semanas_decimales_excedentes = semanas_excedentes - semanas_enteras_excedentes
    incremento = 0
    if semanas_decimales_excedentes > 13 and semanas_decimales_excedentes < 26:
        incremento = 0.5
    if semanas_decimales_excedentes > 26:
        incremento = 1
    total_anios_reconocidos = incremento + anios_excedentes
    
    tabla[11] = tabla[9] * tabla[1] * 365 * total_anios_reconocidos
    tabla[12] = tabla[11] + tabla[10]
    tabla[13] = tabla[12] /12
    
    esposa = tabla[13] * 0.15 * int(tabla[3])
    hijos = tabla[13] * 0.1 * int(tabla[4])
    padres = tabla[13] * 0.10 * int(tabla[5])
    tabla[14] = esposa + hijos + padres #total de asignaciones
    tabla[15] = (tabla[13] + tabla[14]) * 0.11
    tabla[16] = (tabla[13] + tabla[14] + tabla[15])

    #Dependiendo de la edad se agrega la última fila
    EDAD_exacta = relativedelta(datetime.now(), datetime(int(cumple_anio), int(cumple_mes), int(cumple_dia))) #Funcion alterna para sacar meses y dias
    
    #Si se quia ese if, sigue funcionando pero sin redondear edad
    if EDAD <= 60:
        EDAD = 60
    elif EDAD >= 65:
        EDAD = 65
    elif EDAD >60 and EDAD <65:
        if EDAD_exacta.months >=6:
            EDAD = EDAD + 1
        else:
            pass #Edad no cambia

    if EDAD >= 65: 
        tabla[17] = tabla[16] * 1
        porcentaje = 1*100
    if EDAD <= 64: 
        tabla[17] = tabla[16] * 0.95
        porcentaje = 0.95*100
    if EDAD <= 63: 
        tabla[17] = tabla[16] * 0.9
        porcentaje = 0.9*100
    if EDAD <= 62: 
        tabla[17] = tabla[16] * 0.85
        porcentaje = 0.85*100
    if EDAD <= 61: 
        tabla[17] = tabla[16] * 0.8
        porcentaje = 0.8*100
    if EDAD <= 60: 
        tabla[17] = tabla[16] * 0.75
        porcentaje = 0.75*100
    # 60 AÑOS	$3,949	75%
    # 61 AÑOS	$4,212	80%
    # 62 AÑOS	$4,475	85%
    # 63 AÑOS	$4,739	90%
    # 64 AÑOS	$5,002	95%
    # 65 AÑOS	$5,265	100%


    
    #Dando formato final de números
    #Agregando signo de pesos
    tabla[1] = "${:,}".format(tabla[1])
    tabla[8] = round(tabla[8]*100,2)
    tabla[9] = round(tabla[9]*100,2)   
    tabla[10] = "${:,}".format(math.trunc(round(tabla[10],0)))
    tabla[11] = "${:,}".format(math.trunc(round(tabla[11],0)))
    tabla[12] = "${:,}".format(math.trunc(round(tabla[12],0)))
    tabla[13] = "${:,}".format(math.trunc(round(tabla[13],0)))
    tabla[14] = "${:,}".format(math.trunc(round(tabla[14],0)))
    tabla[15] = "${:,}".format(math.trunc(round(tabla[15],0)))
    tabla[16] = "${:,}".format(math.trunc(round(tabla[16],0)))
    tabla[17] = round(tabla[17],0)
    tabla[17] = "${:,}".format(math.trunc(tabla[17]))
    
    # tabla[0]          SEMANAS COTIZADAS
    # tabla[1]          SALARIO PROMEDIO
    # tabla[2]          ASIGNACIONES FAMILIARES
    # tabla[3]          CÓNYUGE
    # tabla[4]          HIJOS (MENORES O ESTUDIANDO)
    # tabla[5]          PADRES (SOLO FALTA DE ESPOSA E HIJOS)
    # tabla[6]          SALARIO MINIMO DF
    # tabla[7]          SALARIO PROMEDIO EN VSM
    # tabla[8]          CUANTÍA BÁSICA
    # tabla[9]          INCREMENTO ANUAL
    # tabla[10]         CUANTÍA BÁSICA ANUAL PESOS
    # tabla[11]         INCREMENTO ANUAL PESOS
    # tabla[12]         ANUALIZADA
    # tabla[13]         MENSUAL
    # tabla[14]         TOTAL ASIGNACIONES
    # tabla[15]         INCREMENTO PRESIDENCIAL    
    # tabla[16]         ****    INCREMENTO PRESIDENCIAL + TOTAL ASIGNACIONES + MENSUAL   ****    
    

    
 
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
    #endregion

    #region Documento

    document = Document("style.docx")
    document._body.clear_content()  #Para limpiar el documento conservando la plantilla
    p = document.add_paragraph('\n\n\n'+date.today().strftime('%d/%m/%Y'))
    p.alignment = 2 # for left, 1 for center, 2 right, 3 justify ....
    p = document.add_paragraph('Nombre: '+NOMBRE)
    p = document.add_paragraph('NSS: '+NSS)
    p = document.add_paragraph('CURP: ' + CURP)
    p = document.add_paragraph('EDAD: ' + str(EDAD_exacta.years) + ' años, '+str(EDAD_exacta.months)+' meses, '+str(EDAD_exacta.days)+' días.')
    p = document.add_paragraph('FECHA NACIMIENTO: ' + str(cumple_dia) + "/" + str(cumple_mes) + "/" + str(cumple_anio) )
    p = document.add_paragraph('UMA 2022: $96,22')
    p = document.add_paragraph('SEMANAS COTIZADAS: '+str(semanas_cotizadas))
    p = document.add_paragraph('SALARIO PROMEDIO: '+str("${:,}".format(salario_prom)))
    if vigente:
        p = document.add_paragraph('ESTATUS: VIGENTE')
    else:
        p = document.add_paragraph('ESTATUS: '+text[60])
        
    p = document.add_paragraph('')


    # SEMANAS COTIZADAS
    # SALARIO PROMEDIO
    # ASIGNACIONES FAMILIARES
    # CÓNYUGE
    # HIJOS (MENORES O ESTUDIANDO)
    # PADRES (SOLO FALTA DE ESPOSA E HIJOS)
    # SALARIO MINIMO DF
    # SALARIO PROMEDIO EN VSM
    # CUANTÍA BÁSICA
    # INCREMENTO ANUAL
    # CUANTÍA BÁSICA ANUAL PESOS
    # INCREMENTO ANUAL PESOS
    # ANUALIZADA
    # MENSUAL
    # TOTAL ASIGNACIONES
    # INCREMENTO PRESIDENCIAL


    tabla_word = document.add_table(1,cols=4)
    tabla_word.rows[0].cells[0].text = "SEMANAS COTIZADAS"
    tabla_word.rows[0].cells[1].text = ""
    tabla_word.rows[0].cells[2].text = "CODIGO"

# CALCULO!C8<=800,"2",SI(CALCULO!C8<1200,"3",SI(CALCULO!C8<1600,"4",SI(CALCULO!C8>1600,"5"))))
    semanas_cotizadas = int(semanas_cotizadas)

    # input(vigente)
    # input(semanas_cotizadas)
    print(vigente)
    print(semanas_cotizadas)

    if vigente:
        if semanas_cotizadas <= 800:
            color = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
            codigo="VVD6"
        elif semanas_cotizadas < 1200:
            color = parse_xml(r'<w:shd {} w:fill="FF9900"/>'.format(nsdecls('w')))
            codigo = "VVC7"
        elif semanas_cotizadas < 1600:
            color = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
            codigo = "VVB8"
        elif semanas_cotizadas >= 1600:
            color = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
            codigo = "VVA10"
        tabla_word.rows[0].cells[3]._tc.get_or_add_tcPr().append(color)
        tabla_word.rows[0].cells[3].text = codigo
    else: 
        anio = diferencia_fechas(text[60],date.today().strftime('%d/%m/%Y')) 
        anio = int(round(anio/365,0))

        if anio < 1: anio = 1 #Si apenas está por cumplir el año
        print(anio)
        codigo = "RR"
        if semanas_cotizadas <= 800:
            codigo = codigo + "D"
        elif semanas_cotizadas < 1200:
            codigo = codigo + "C"
        elif semanas_cotizadas < 1600:
            codigo = codigo + "B"
        elif semanas_cotizadas >= 1600:
            codigo = codigo + "A"

        if anio == 1: codigo = codigo + "6"
        if anio == 2: codigo = codigo + "7"
        if anio == 3: codigo = codigo + "8"
        if anio == 4: codigo = codigo + "9"
        if anio == 5: codigo = codigo + "10"
        # input(codigo)
        verde_fuerte = ['RRA6', 'RRA7', 'RRA8', 'RRA9', 'RRA10', 'RRB9', 'RRB10', 'RRC10']
        verde_claro = ['RRB7', 'RRB8', 'RRC8', 'RRC9']
        amarillo = ['RRB6', 'RRC7']
        naranja = ['RRD10','RRC6']
        rojo = ['RRD6','RRD7','RRD8','RRD9']
        if codigo in rojo:
            color = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
        if codigo in naranja:
            color = parse_xml(r'<w:shd {} w:fill="FF9900"/>'.format(nsdecls('w')))
        if codigo in amarillo:
            color = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
        if codigo in verde_claro:
            color = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
        if codigo in verde_fuerte:
            color = parse_xml(r'<w:shd {} w:fill="00B050"/>'.format(nsdecls('w')))

        if anio < 6:
            tabla_word.rows[0].cells[3]._tc.get_or_add_tcPr().append(color)
            tabla_word.rows[0].cells[3].text = codigo
        else:
            tabla_word.rows[0].cells[3].text = "No aplica"









    for i in range(17): tabla_word.add_row().cells
    
    filas = []
    filas.append("SEMANAS COTIZADAS")
    filas.append("SALARIO PROMEDIO")
    filas.append("ASIGNACIONES FAMILIARES")
    filas.append("CÓNYUGE")
    filas.append("HIJOS (MENORES O ESTUDIANDO)")
    filas.append("PADRES (SOLO FALTA DE ESPOSA E HIJOS)")
    filas.append("SALARIO MINIMO DF")
    filas.append("SALARIO PROMEDIO EN VSM")
    filas.append("CUANTÍA BÁSICA")
    filas.append("INCREMENTO ANUAL")
    filas.append("CUANTÍA BÁSICA ANUAL PESOS")
    filas.append("INCREMENTO ANUAL PESOS")
    filas.append("ANUALIZADA")
    filas.append("MENSUAL")
    filas.append("TOTAL ASIGNACIONES")
    filas.append("INCREMENTO PRESIDENCIAL")
    filas.append("")
    filas.append(str(EDAD) + " AÑOS.")
    
    # Llenando contenido
    for i, fila in enumerate(filas):
        tabla_word.rows[i].cells[0].text = fila                             #Escribir el elemento de filas en el renglon
        if i<len(tabla): tabla_word.rows[i].cells[1].text = str(tabla[i])   #Escribir los datos numericos de tabla
    tabla_word.rows[16].cells[2].text = "100%"
    tabla_word.rows[17].cells[2].text = str(round(porcentaje,0)) + "%"
    # tabla_word.style = "Table Grid"
    
    
    for row in tabla_word.rows:
        for i, cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            if i == 2:                  #Centrar columna 2
                paragraph.alignment = 1 #Centrar columna 2
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(9)
                    font.color.rgb = RGBColor(0, 0, 0)


                    
                    
    #Color a la ultima fila
    for i in range(3):
        color = parse_xml(r'<w:shd {} w:fill="CC00FF"/>'.format(nsdecls('w')))
        tabla_word.rows[-1].cells[i]._tc.get_or_add_tcPr().append(color)
    
    #Ancho a fila
    tabla_word.autofit = False 
    tabla_word.allow_autofit = False
    widths = (Cm(7.34), Cm(6.17),Cm(4.87),Cm(2.58))
    for row in tabla_word.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


    


    p = document.add_paragraph().add_run('')
    p = document.add_paragraph().add_run('OBSERVACIONES')
    font = p.font
    font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)


    p = document.add_paragraph('En la consulta se consideraron las semanas al día ' +date.today().strftime('%d/%m/%Y') + " y el salario promedio de la misma constancia. ")
    
    r = p.add_run("Representa la pensión que la persona recibiría hoy ")
    r.bold = True
    r.underline = True


    if EDAD < 60:
        r = p.add_run("al cumplir 60 años. ")
        r.bold = True
        r.underline = True
    else: p.add_run("")


    p.add_run("En la constancia electrónica se reconocen "+ str(semanas_reconocidas) +" semanas. "+ str(semanas_desconocidas) +" semanas descontadas. "+str(semanas_reintegradas)+" semanas reintegradas. ")
    

    p.add_run("Recibiría pensión mínima garantizada que en 2022 es de $"+str(pension_minima)+" pesos.").font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    p.alignment=3
    p = document.add_paragraph("\n")
    p.add_run("Le podemos hacer una propuesta para ayudarlo a obtener una pensión arriba de ")
    
    r = p.add_run("**"    +" mil pesos ")
    r.bold = True
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p.add_run("desde ")

    p.add_run("** de 202*. ").bold = True

    p.add_run("Hagamos una cita para ver si califica. ")
    p.add_run("(+"+str(int(semanas_desconocidas)-int(semanas_reintegradas))+").")
    p.add_run(" Pudiera haber semanas no reconocidas (19**-19**)")
    
    p = document.add_paragraph("\n*El monto de su pensión puede llegar a variar dependiendo de su baja con el patrón. ")
    if vigente:
        p.add_run("Baja esperada al día "+ultimo_dia_mes+" ")

    
    palabras =text[6].split(' ')
    #Para cambiar orden del nombre, y nombrar así el documento
    compuestos = ['da', 'de', 'del', 'la', 'las', 'los','y', 'i', 'san', 'santa']
    if compuestos in palabras or len(palabras) > 4:
        array_document_name = "demo_"
    elif len(palabras) == 3:
        array_document_name = palabras[-1:] + palabras[:-1] 
    elif len(palabras) == 4:
        array_document_name = palabras[-2:] + palabras[:-2] 
    document_name = ""
    for i in array_document_name:
        document_name = document_name+i+" "
    
    document_name = r"DIAGNOSTICOS/demo_" +document_name[:-1].replace(" ","_") +".docx"
    document.save(document_name)
    #endregion

    os.system("start " + document_name)



    

if __name__ == "__main__":
    os.system("taskkill /IM WINWORD.exe")
    os.system('cls')
    diagnostico(r"PDF\constancia_CARLOS MARTIN ANGUIANO GARCIA.pdf")
    diagnostico(r"PDF\test0.pdf")
    # diagnostico(r"PDF\Constancia FELIPE ZAMARRIPA CERVANTES.pdf")
    # diagnostico(r"PDF\constancia_JUANA MARIA CARREON RODRIGUEZ.pdf")
